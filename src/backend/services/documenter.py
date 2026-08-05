import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.backend.services.db_manager import save_report_entry

# ---------------------------------------------------------------------------
# Colour palette – everything visible on white paper
# ---------------------------------------------------------------------------
BLACK      = RGBColor(0x00, 0x00, 0x00)   # body text, table cells
DARK_GREY  = RGBColor(0x33, 0x33, 0x33)   # footer / captions
MID_GREY   = RGBColor(0x55, 0x55, 0x55)   # subtitle line
ACCENT     = RGBColor(0x1A, 0x1A, 0x2E)   # title block (near-black navy)

# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _add_page_number(run):
    """Inserts a PAGE field into a docx Run element."""
    r = run._r
    for tag, ftype in [('w:fldChar', 'begin'), ('w:fldChar', 'separate'), ('w:fldChar', 'end')]:
        elem = OxmlElement(tag)
        if ftype in ('begin', 'end', 'separate'):
            elem.set(qn('w:fldCharType'), ftype)
        r.append(elem)
        if ftype == 'begin':
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = ' PAGE '
            r.append(instr)


def _add_toc_field(paragraph):
    """Inserts a dynamic TOC field into a Paragraph."""
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-3" \h \z \u '
    r.append(instr)

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r.append(fld_sep)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r.append(fld_end)


def _apply_run_font(run, size_pt=11, bold=False, color=None, name='Courier New'):
    """Centralised run styling that always enforces a visible colour."""
    run.font.name  = name
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.color.rgb = color if color is not None else BLACK


def _set_cell_text(cell, text, bold=False, size_pt=11):
    """Clears a table cell and writes text with enforced black font."""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    _apply_run_font(run, size_pt=size_pt, bold=bold, color=BLACK)


# ---------------------------------------------------------------------------
# Main service class
# ---------------------------------------------------------------------------

class MeetingDocumenter:
    """Generates structured, aesthetically professional Word reports and
    archives metadata in the SQLite database via db_manager."""

    # __file__ lives at:  src/backend/services/documenter.py
    # We need 4 dirname() calls to reach the project root.
    _ROOT_DIR = os.path.dirname(
        os.path.dirname(
            os.path.dirname(
                os.path.dirname(os.path.abspath(__file__))
            )
        )
    )

    def _reports_dir(self):
        path = os.path.join(self._ROOT_DIR, "reports")
        os.makedirs(path, exist_ok=True)
        return path

    # ------------------------------------------------------------------

    def generar_reporte(self, analisis_data):
        """Generates an aesthetic .docx report and logs metadata to the DB.

        Args:
            analisis_data (dict): Must contain:
                - ``sentiment``      (float)
                - ``subjectivity``   (float)
                - ``total_phrases``  (int)
                - ``top_keywords``   (list[str])
                - ``texto_original`` (str) – **must be** ``' '.join(recording_buffer)``
                  i.e. the full stitched session transcript assembled by the
                  BackendController before calling this method.  Passing only the
                  last chunk here will produce an incomplete document.

        Returns:
            str: Absolute filepath of the saved .docx file.
        """
        doc = Document()

        # ----------------------------------------------------------------
        # 1. Global Normal style – black Courier New, 11 pt
        # ----------------------------------------------------------------
        normal_style = doc.styles['Normal']
        nf = normal_style.font
        nf.name  = 'Courier New'
        nf.size  = Pt(11)
        nf.color.rgb = BLACK

        # Heading styles – black, bold, Courier New
        for level, size in ((1, 14), (2, 12), (3, 11)):
            h_style = doc.styles[f'Heading {level}']
            hf = h_style.font
            hf.name  = 'Courier New'
            hf.size  = Pt(size)
            hf.bold  = True
            hf.color.rgb = ACCENT if level == 1 else BLACK

        # ----------------------------------------------------------------
        # 2. Page setup & footer
        # ----------------------------------------------------------------
        for section in doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.25)
            section.right_margin  = Inches(1.25)

            footer = section.footer
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp.clear()   # wipe any default empty run

            brand_run = fp.add_run(
                "RRDOCS  //  SINAPSIS BREACH PROTOCOL  //  Confidential  |  Page "
            )
            _apply_run_font(brand_run, size_pt=8, color=DARK_GREY)

            page_run = fp.add_run()
            _apply_run_font(page_run, size_pt=8, color=DARK_GREY)
            _add_page_number(page_run)

        # ----------------------------------------------------------------
        # 3. Title block  (level=0 → built-in Title style)
        # ----------------------------------------------------------------
        title_p = doc.add_heading("SYS.RRDOCS // EXECUTIVE BREACH REPORT", level=0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_p.runs:
            _apply_run_font(run, size_pt=18, bold=True, color=ACCENT, name='Courier New')

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_str = datetime.now().strftime("%Y-%m-%d  %H:%M:%S  UTC")
        sub_run = sub_p.add_run(f"GENERATED: {timestamp_str}")
        _apply_run_font(sub_run, size_pt=10, color=MID_GREY)

        doc.add_paragraph()  # spacer

        # ----------------------------------------------------------------
        # 4. Table of Contents
        # ----------------------------------------------------------------
        toc_heading = doc.add_heading("TABLE OF CONTENTS", level=2)
        for run in toc_heading.runs:
            _apply_run_font(run, size_pt=12, bold=True, color=BLACK)

        toc_p = doc.add_paragraph()
        _add_toc_field(toc_p)

        doc.add_page_break()

        # ----------------------------------------------------------------
        # 5. Section 1 – Telemetry & Sentiment
        # ----------------------------------------------------------------
        doc.add_heading("1.  TELEMETRY & SENTIMENT", level=1)

        metrics = [
            ("Sentiment Polarity",  f"{analisis_data.get('sentiment', 0.0):+.4f}"),
            ("Subjectivity Score",  f"{analisis_data.get('subjectivity', 0.0):.4f}"),
            ("Total Phrase Count",  str(analisis_data.get('total_phrases', 0))),
        ]
        for label, value in metrics:
            p = doc.add_paragraph(style='Normal')
            label_run = p.add_run(f"{label:<22}")
            _apply_run_font(label_run, size_pt=11, bold=True, color=BLACK)
            val_run = p.add_run(f":  {value}")
            _apply_run_font(val_run, size_pt=11, color=BLACK)

        doc.add_paragraph()  # spacer

        # ----------------------------------------------------------------
        # 6. Section 2 – Executive Summary  (LLM generated)
        # ----------------------------------------------------------------
        doc.add_heading("2.  EXECUTIVE SUMMARY", level=1)

        resumen = analisis_data.get("resumen_ejecutivo", [])
        if resumen:
            for i, punto in enumerate(resumen, 1):
                p = doc.add_paragraph(style='Normal')
                bullet_run = p.add_run(f"  {i}.  ")
                _apply_run_font(bullet_run, size_pt=11, bold=True, color=BLACK)
                text_run = p.add_run(str(punto))
                _apply_run_font(text_run, size_pt=11, color=BLACK)
        else:
            p = doc.add_paragraph(style='Normal')
            _apply_run_font(p.add_run("[No executive summary generated]"), color=BLACK)

        doc.add_paragraph()  # spacer

        # ----------------------------------------------------------------
        # 7. Section 3 – Top Concepts Table  (LLM generated)
        # ----------------------------------------------------------------
        doc.add_heading("3.  KEY CONCEPTS", level=1)

        intro_p = doc.add_paragraph(style='Normal')
        _apply_run_font(
            intro_p.add_run(
                "Top technical and philosophical concepts identified by the AI model:"
            ),
            color=BLACK,
        )
        doc.add_paragraph()

        conceptos = analisis_data.get("conceptos_principales", [])
        # Fall back to frequency-based keywords if LLM produced nothing
        if not conceptos:
            conceptos = analisis_data.get("top_keywords", [])

        # Compute word frequencies from raw transcript for frequency column
        texto_original = analisis_data.get("texto_original", "")
        words = re.findall(r"[a-zA-ZáéíóúÁÉÍÓÚñÑüÜ]+", texto_original.lower())
        freq_map = {}
        for w in words:
            freq_map[w] = freq_map.get(w, 0) + 1

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        _set_cell_text(hdr[0], "CONCEPT",   bold=True, size_pt=11)
        _set_cell_text(hdr[1], "FREQUENCY", bold=True, size_pt=11)
        for kw in conceptos:
            row = table.add_row().cells
            _set_cell_text(row[0], kw.upper(), size_pt=11)
            _set_cell_text(row[1], str(freq_map.get(kw.lower(), "—")), size_pt=11)
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph()  # spacer

        # ----------------------------------------------------------------
        # 8. Section 4 – Cognitive Biases  (LLM generated)
        # ----------------------------------------------------------------
        doc.add_heading("4.  COGNITIVE BIAS ANALYSIS", level=1)

        biases = analisis_data.get("sesgos_cognitivos", [])
        if biases:
            for sesgo in biases:
                p = doc.add_paragraph(style='Normal')
                bullet = p.add_run("  •  ")
                _apply_run_font(bullet, size_pt=11, bold=True, color=BLACK)
                text_r = p.add_run(str(sesgo))
                _apply_run_font(text_r, size_pt=11, color=BLACK)
        else:
            p = doc.add_paragraph(style='Normal')
            _apply_run_font(
                p.add_run("[No cognitive biases identified]"),
                color=BLACK,
            )

        doc.add_paragraph()  # spacer

        # ----------------------------------------------------------------
        # 10. Section 5 – Action Items (LLM generated)
        # ----------------------------------------------------------------
        doc.add_heading("5.  ACTION ITEMS", level=1)

        tareas = analisis_data.get("tareas_acciones", [])
        if tareas:
            for tarea in tareas:
                p = doc.add_paragraph(style='Normal')
                bullet = p.add_run("  [ ]  ")
                _apply_run_font(bullet, size_pt=11, bold=True, color=BLACK)
                text_r = p.add_run(str(tarea))
                _apply_run_font(text_r, size_pt=11, color=BLACK)
        else:
            p = doc.add_paragraph(style='Normal')
            _apply_run_font(
                p.add_run("[No action items identified]"),
                color=BLACK,
            )

        doc.add_paragraph()  # spacer

        # ----------------------------------------------------------------
        # 11. Section 6 – Corrected Transcript (LLM generated)
        # ----------------------------------------------------------------
        doc.add_heading("6.  IA CORRECTED TRANSCRIPT", level=1)
        
        texto_corregido = analisis_data.get("texto_corregido", [])
        if texto_corregido:
            for parrafo in texto_corregido:
                p = doc.add_paragraph(style='Normal')
                _apply_run_font(p.add_run(str(parrafo)), size_pt=11, color=BLACK)
        else:
            p = doc.add_paragraph(style='Normal')
            _apply_run_font(p.add_run("[No corrected transcript available]"), color=BLACK)

        doc.add_paragraph()  # spacer

        # ----------------------------------------------------------------
        # 12. Section 7 – Raw Transcript
        # ----------------------------------------------------------------
        doc.add_heading("7.  RAW TRANSCRIPT", level=1)

        transcript_p = doc.add_paragraph(style='Normal')
        transcript_run = transcript_p.add_run(
            texto_original if texto_original else "[No transcript available]"
        )
        _apply_run_font(transcript_run, size_pt=11, color=BLACK)

        # ----------------------------------------------------------------
        # 13. Save document
        # ----------------------------------------------------------------
        file_ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"reporte_reunion_{file_ts}.docx"
        filepath = os.path.join(self._reports_dir(), filename)
        doc.save(filepath)
        print(f"[Documenter] Report saved -> {filepath}")

        # ----------------------------------------------------------------
        # 14. Archive metadata in database
        # ----------------------------------------------------------------
        db_entry = {
            "sentiment_score": analisis_data.get("sentiment", 0.0),
            # Store all LLM-extracted fields in the keywords JSON column
            "keywords": {
                "conceptos":  analisis_data.get("conceptos_principales", []),
                "resumen":    analisis_data.get("resumen_ejecutivo",     []),
                "sesgos":     analisis_data.get("sesgos_cognitivos",     []),
                "tareas":     analisis_data.get("tareas_acciones",       []),
            },
            "file_path":  filepath,
            "timestamp":  datetime.now().isoformat(),
        }
        save_report_entry(db_entry)

        return filepath
